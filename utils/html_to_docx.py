"""HTML → .docx converter for editor-authored documents (NOC letters).

Produces a **real OOXML file** with python-docx, so the download opens in Word /
LibreOffice fully editable — not an HTML file wearing a .doc extension.

This is a **bounded** converter, deliberately. It covers exactly the shapes our
NOC templates and CKEditor's toolbar produce:

    p div h1-h6 br hr span strong/b em/i u s strike sub sup a code mark
    ul ol li  table thead/tbody/tr/th/td  img

plus ``text-align``, ``font-size``, ``font-weight``, ``font-style`` and
``text-decoration`` from inline styles. Anything else degrades to plain
paragraphs with its text intact rather than raising — an operator pasting odd
markup from Word should never get a 500.

**Bangla in Word**: python-docx's ``run.font.name`` only sets ``w:ascii`` and
``w:hAnsi``. Bengali is a *complex script*, so Word picks the font from
``w:cs`` — without it Bangla falls back to Times New Roman and renders as
boxes. ``_set_run_font()`` sets all three (plus ``w:eastAsia``).
"""

import base64
import io
import os
import re
from html.parser import HTMLParser

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor

DEFAULT_FONT = 'SolaimanLipi'
DEFAULT_SIZE_PT = 12.0

VOID_TAGS = {'br', 'hr', 'img', 'col', 'meta', 'link', 'input'}

BLOCK_TAGS = {
    'p', 'div', 'section', 'blockquote', 'pre', 'li', 'hr',
    'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
    'table', 'thead', 'tbody', 'tfoot', 'tr', 'td', 'th', 'ul', 'ol',
    'figure', 'figcaption', 'caption',
}

HEADING_SIZES = {'h1': 18.0, 'h2': 16.0, 'h3': 14.0, 'h4': 13.0, 'h5': 12.0, 'h6': 11.0}

ALIGN_MAP = {
    'left': WD_ALIGN_PARAGRAPH.LEFT,
    'right': WD_ALIGN_PARAGRAPH.RIGHT,
    'center': WD_ALIGN_PARAGRAPH.CENTER,
    'centre': WD_ALIGN_PARAGRAPH.CENTER,
    'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
}

_WS_RE = re.compile(r'[ \t\r\n]+')


# ── tiny DOM ─────────────────────────────────────────────────────────────────

class _Node:
    __slots__ = ('tag', 'attrs', 'children', 'text')

    def __init__(self, tag, attrs=None, text=None):
        self.tag = tag
        self.attrs = dict(attrs or {})
        self.children = []
        self.text = text

    @property
    def is_text(self):
        return self.tag is None

    def style(self):
        """Parsed inline ``style`` declarations as a lowercase dict."""
        out = {}
        for decl in (self.attrs.get('style') or '').split(';'):
            if ':' in decl:
                prop, _, val = decl.partition(':')
                out[prop.strip().lower()] = val.strip()
        return out


class _TreeBuilder(HTMLParser):
    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.root = _Node('root')
        self.stack = [self.root]

    def handle_starttag(self, tag, attrs):
        tag = tag.lower()
        node = _Node(tag, {k.lower(): (v or '') for k, v in attrs})
        self.stack[-1].children.append(node)
        if tag not in VOID_TAGS:
            self.stack.append(node)

    def handle_startendtag(self, tag, attrs):
        tag = tag.lower()
        node = _Node(tag, {k.lower(): (v or '') for k, v in attrs})
        self.stack[-1].children.append(node)

    def handle_endtag(self, tag):
        tag = tag.lower()
        if tag in VOID_TAGS:
            return
        # Pop to the matching open tag; tolerate stray/unbalanced end tags.
        for i in range(len(self.stack) - 1, 0, -1):
            if self.stack[i].tag == tag:
                del self.stack[i:]
                return

    def handle_data(self, data):
        self.stack[-1].children.append(_Node(None, text=data))


# ── inline formatting state ──────────────────────────────────────────────────

class _Fmt:
    __slots__ = ('bold', 'italic', 'underline', 'strike', 'size', 'sub', 'sup', 'color')

    def __init__(self, bold=False, italic=False, underline=False, strike=False,
                 size=None, sub=False, sup=False, color=None):
        self.bold, self.italic, self.underline, self.strike = bold, italic, underline, strike
        self.size, self.sub, self.sup, self.color = size, sub, sup, color

    def copy(self):
        return _Fmt(self.bold, self.italic, self.underline, self.strike,
                    self.size, self.sub, self.sup, self.color)


def _parse_size(value):
    """CSS font-size → points. Handles pt, px, em/rem and bare numbers."""
    if not value:
        return None
    m = re.match(r'^\s*([\d.]+)\s*(pt|px|em|rem|%)?\s*$', value)
    if not m:
        return None
    num = float(m.group(1))
    unit = m.group(2) or 'px'
    if unit == 'pt':
        return num
    if unit == 'px':
        return num * 0.75                     # 96dpi CSS px → pt
    if unit in ('em', 'rem'):
        return num * DEFAULT_SIZE_PT
    if unit == '%':
        return DEFAULT_SIZE_PT * num / 100.0
    return None


def _parse_color(value):
    if not value:
        return None
    v = value.strip()
    m = re.match(r'^#([0-9a-fA-F]{6})$', v)
    if m:
        return RGBColor.from_string(m.group(1).upper())
    m = re.match(r'^#([0-9a-fA-F]{3})$', v)
    if m:
        d = m.group(1)
        return RGBColor.from_string(''.join(c * 2 for c in d).upper())
    m = re.match(r'^rgb\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*\)$', v)
    if m:
        return RGBColor(*(int(g) for g in m.groups()))
    return None


def _apply_tag_format(node, fmt):
    """Fold one element's own emphasis + inline style into a copy of `fmt`."""
    fmt = fmt.copy()
    tag = node.tag
    if tag in ('strong', 'b'):
        fmt.bold = True
    elif tag in ('em', 'i'):
        fmt.italic = True
    elif tag == 'u':
        fmt.underline = True
    elif tag in ('s', 'strike', 'del'):
        fmt.strike = True
    elif tag == 'sub':
        fmt.sub = True
    elif tag == 'sup':
        fmt.sup = True
    elif tag in HEADING_SIZES:
        fmt.bold = True
        fmt.size = HEADING_SIZES[tag]
    elif tag == 'th':
        fmt.bold = True

    css = node.style()
    weight = css.get('font-weight', '')
    if weight in ('bold', 'bolder') or (weight.isdigit() and int(weight) >= 600):
        fmt.bold = True
    elif weight in ('normal', '400'):
        fmt.bold = False
    if css.get('font-style') in ('italic', 'oblique'):
        fmt.italic = True
    decoration = css.get('text-decoration', '') + ' ' + css.get('text-decoration-line', '')
    if 'underline' in decoration:
        fmt.underline = True
    if 'line-through' in decoration:
        fmt.strike = True
    size = _parse_size(css.get('font-size'))
    if size:
        fmt.size = size
    color = _parse_color(css.get('color'))
    if color is not None:
        fmt.color = color
    return fmt


def _find_align(node):
    css = node.style()
    align = (css.get('text-align') or node.attrs.get('align') or '').strip().lower()
    return ALIGN_MAP.get(align)


# ── docx helpers ─────────────────────────────────────────────────────────────

def _set_run_font(run, font_name, size_pt):
    """Set the run's font for latin AND complex scripts.

    ``w:cs`` is the one that matters for Bangla — Word selects the complex-script
    font from it, and without it Bengali text falls back to the theme font.
    """
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.insert(0, rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rfonts.set(qn(attr), font_name)
    # Complex-script size is a separate property; without it Bangla ignores w:sz.
    szcs = rpr.find(qn('w:szCs'))
    if szcs is None:
        szcs = rpr.makeelement(qn('w:szCs'), {})
        rpr.append(szcs)
    szcs.set(qn('w:val'), str(int(round(size_pt * 2))))


def _set_table_borders(table, on):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.makeelement(qn('w:tblBorders'), {})
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = borders.makeelement(qn(f'w:{edge}'), {})
        el.set(qn('w:val'), 'single' if on else 'none')
        el.set(qn('w:sz'), '4' if on else '0')
        el.set(qn('w:color'), '000000' if on else 'auto')
        borders.append(el)
    tbl_pr.append(borders)


def _wants_borders(node):
    """True when the HTML asks for visible table rules."""
    border_attr = (node.attrs.get('border') or '').strip()
    if border_attr and border_attr != '0':
        return True
    css_border = node.style().get('border', '')
    return bool(css_border) and 'none' not in css_border and '0' not in css_border.split()[:1]


def _resolve_image(src, static_root):
    """`src` (data: URI or /static/ path) → a BytesIO python-docx can embed."""
    if not src:
        return None
    if src.startswith('data:image/'):
        _, _, payload = src.partition(',')
        try:
            return io.BytesIO(base64.b64decode(payload))
        except Exception:
            return None
    if not static_root:
        return None
    # Map "/static/img/logo.png" (or "img/logo.png") onto the source static dir.
    rel = src.split('?')[0].split('#')[0]
    for prefix in ('/static/', 'static/', '/'):
        if rel.startswith(prefix):
            rel = rel[len(prefix):]
            break
    path = os.path.normpath(os.path.join(static_root, rel))
    if not path.startswith(os.path.normpath(static_root)) or not os.path.exists(path):
        return None
    with open(path, 'rb') as fh:
        return io.BytesIO(fh.read())


# ── the renderer ─────────────────────────────────────────────────────────────

class _Renderer:
    def __init__(self, doc, font, base_size, static_root):
        self.doc = doc
        self.font = font
        self.base_size = base_size
        self.static_root = static_root

    # -- inline ---------------------------------------------------------------
    def _add_text(self, paragraph, text, fmt):
        # Collapse HTML whitespace, but keep a single separating space.
        text = _WS_RE.sub(' ', text)
        if not text:
            return
        run = paragraph.add_run(text)
        run.bold = fmt.bold
        run.italic = fmt.italic
        run.underline = fmt.underline
        if fmt.strike:
            run.font.strike = True
        if fmt.sub:
            run.font.subscript = True
        if fmt.sup:
            run.font.superscript = True
        if fmt.color is not None:
            run.font.color.rgb = fmt.color
        _set_run_font(run, self.font, fmt.size or self.base_size)

    def _add_image(self, paragraph, node):
        stream = _resolve_image(node.attrs.get('src'), self.static_root)
        if stream is None:
            alt = node.attrs.get('alt')
            if alt:
                self._add_text(paragraph, alt, _Fmt(italic=True))
            return
        width = None
        raw_w = node.attrs.get('width') or _parse_size(node.style().get('width'))
        try:
            if raw_w:
                px = float(str(raw_w).replace('px', '').strip())
                width = Inches(px / 96.0)
        except (TypeError, ValueError):
            width = None
        try:
            paragraph.add_run().add_picture(stream, width=width or Inches(0.9))
        except Exception:
            pass        # an unreadable image must not sink the whole export

    def render_inline(self, node, paragraph, fmt):
        for child in node.children:
            if child.is_text:
                self._add_text(paragraph, child.text, fmt)
            elif child.tag == 'br':
                paragraph.add_run().add_break()
            elif child.tag == 'img':
                self._add_image(paragraph, child)
            else:
                self.render_inline(child, paragraph, _apply_tag_format(child, fmt))

    # -- blocks ---------------------------------------------------------------
    def _new_paragraph(self, node, align, style=None):
        p = self.doc.add_paragraph(style=style) if style else self.doc.add_paragraph()
        node_align = _find_align(node) if node is not None else None
        if node_align is not None:
            p.alignment = node_align
        elif align is not None:
            p.alignment = align
        p.paragraph_format.space_after = Pt(4)
        return p

    def render_children(self, node, fmt, align):
        """Walk `node`'s children, batching runs of inline content into paragraphs."""
        pending = []

        def flush():
            if not pending:
                return
            if all(c.is_text and not c.text.strip() for c in pending):
                pending.clear()
                return
            paragraph = self._new_paragraph(node, align)
            holder = _Node('span')
            holder.children = list(pending)
            self.render_inline(holder, paragraph, fmt)
            pending.clear()

        for child in node.children:
            if child.is_text or child.tag not in BLOCK_TAGS:
                pending.append(child)
            else:
                flush()
                self.render_block(child, fmt, align)
        flush()

    def render_block(self, node, fmt, align):
        tag = node.tag
        fmt = _apply_tag_format(node, fmt)
        node_align = _find_align(node)
        align = node_align if node_align is not None else align

        if tag == 'hr':
            p = self.doc.add_paragraph()
            pbdr = p._p.get_or_add_pPr().makeelement(qn('w:pBdr'), {})
            bottom = pbdr.makeelement(qn('w:bottom'), {})
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:color'), '000000')
            pbdr.append(bottom)
            p._p.get_or_add_pPr().append(pbdr)
            return

        if tag == 'table':
            self.render_table(node, fmt, align)
            return

        if tag in ('ul', 'ol'):
            style = 'List Bullet' if tag == 'ul' else 'List Number'
            for item in node.children:
                if item.is_text or item.tag != 'li':
                    continue
                p = self._new_paragraph(item, align, style=style)
                self.render_inline(item, p, _apply_tag_format(item, fmt))
                # A nested list inside the <li> still needs rendering.
                for sub in item.children:
                    if not sub.is_text and sub.tag in ('ul', 'ol', 'table'):
                        self.render_block(sub, fmt, align)
            return

        if tag in ('div', 'section', 'figure', 'tbody', 'thead', 'tfoot'):
            self.render_children(node, fmt, align)
            return

        # p, h1-h6, blockquote, pre, li, caption, figcaption — one paragraph,
        # unless it contains block children, in which case recurse.
        has_block_child = any(
            (not c.is_text) and c.tag in BLOCK_TAGS for c in node.children)
        if has_block_child:
            self.render_children(node, fmt, align)
            return

        p = self._new_paragraph(node, align)
        if tag in HEADING_SIZES:
            p.paragraph_format.space_before = Pt(6)
        self.render_inline(node, p, fmt)

    def render_table(self, node, fmt, align):
        rows = []
        for tr in self._iter_rows(node):
            cells = [c for c in tr.children if not c.is_text and c.tag in ('td', 'th')]
            if cells:
                rows.append(cells)
        if not rows:
            return

        n_cols = max(sum(self._colspan(c) for c in row) for row in rows)
        table = self.doc.add_table(rows=len(rows), cols=n_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        _set_table_borders(table, _wants_borders(node))

        for r, row in enumerate(rows):
            col = 0
            for cell_node in row:
                if col >= n_cols:
                    break
                span = min(self._colspan(cell_node), n_cols - col)
                cell = table.cell(r, col)
                if span > 1:
                    cell = cell.merge(table.cell(r, col + span - 1))
                # The cell arrives with one empty paragraph — render into it.
                cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
                saved = self.doc
                self.doc = cell
                try:
                    cell_align = _find_align(cell_node) or align
                    self.render_block(cell_node, fmt, cell_align)
                finally:
                    self.doc = saved
                if not cell.paragraphs:
                    cell.add_paragraph()
                col += span

    @staticmethod
    def _colspan(cell_node):
        try:
            return max(1, int(cell_node.attrs.get('colspan', '1')))
        except (TypeError, ValueError):
            return 1

    def _iter_rows(self, node):
        for child in node.children:
            if child.is_text:
                continue
            if child.tag == 'tr':
                yield child
            elif child.tag in ('thead', 'tbody', 'tfoot'):
                yield from self._iter_rows(child)


# ── public API ───────────────────────────────────────────────────────────────

def html_to_docx(html, *, font=DEFAULT_FONT, base_size=DEFAULT_SIZE_PT,
                 static_root=None, margins_mm=(15, 18, 15, 18), title=None):
    """Convert `html` to a .docx and return it as ``BytesIO`` (seeked to 0).

    ``margins_mm`` is (top, right, bottom, left). ``static_root`` is the
    filesystem directory ``/static/...`` image sources resolve against.
    """
    doc = Document()

    section = doc.sections[0]
    section.page_width = Mm(210)          # A4 portrait
    section.page_height = Mm(297)
    top, right, bottom, left = margins_mm
    section.top_margin = Mm(top)
    section.right_margin = Mm(right)
    section.bottom_margin = Mm(bottom)
    section.left_margin = Mm(left)

    # Make the document default match, so text typed in Word after download
    # keeps the same font instead of reverting to Calibri.
    normal = doc.styles['Normal']
    normal.font.name = font
    normal.font.size = Pt(base_size)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.insert(0, rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rfonts.set(qn(attr), font)

    if title:
        doc.core_properties.title = title

    builder = _TreeBuilder()
    builder.feed(html or '')
    builder.close()

    renderer = _Renderer(doc, font, base_size, static_root)
    renderer.render_children(builder.root, _Fmt(), None)

    # A brand-new Document starts with one empty paragraph; drop it if the
    # content produced anything of its own.
    body = doc.element.body
    paragraphs = doc.paragraphs
    if len(paragraphs) > 1 and not paragraphs[0].text.strip() and not paragraphs[0].runs:
        body.remove(paragraphs[0]._p)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
